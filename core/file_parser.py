from PyPDF2 import PdfReader
from docx import Document
from fastapi import UploadFile
from PIL import Image
import pytesseract
import io
import fitz  # PyMuPDF

# ✅ Path Tesseract (ubah sesuai environment)
# pytesseract.pytesseract.tesseract_cmd = r"D:\App\Tesseract-OCR\tesseract.exe"


# ==========================
# 🔹 EKSTRAK TEKS DARI FILE
# ==========================
def extract_text(file: UploadFile):
    filename = file.filename.lower()
    file.file.seek(0)

    # === PDF ===
    if filename.endswith(".pdf"):
        reader = PdfReader(file.file)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text.strip()

    # === DOCX ===
    elif filename.endswith(".docx"):
        file.file.seek(0)
        doc = Document(io.BytesIO(file.file.read()))
        texts = []

        # Ambil teks dari paragraf biasa
        for para in doc.paragraphs:
            if para.text.strip():
                texts.append(para.text.strip())

        # Ambil teks dari tabel
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = "\n".join(
                        p.text.strip() for p in cell.paragraphs if p.text.strip()
                    )
                    if cell_text:
                        texts.append(cell_text)

        # Ambil teks dari gambar dalam dokumen (OCR)
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                image_data = rel.target_part.blob
                image = Image.open(io.BytesIO(image_data)).convert("RGB")
                ocr_text = pytesseract.image_to_string(image, lang="eng+ind").strip()
                if ocr_text:
                    texts.append(ocr_text)

        return "\n".join(texts).strip()

    # === TXT ===
    elif filename.endswith(".txt"):
        return file.file.read().decode("utf-8", errors="ignore")

    else:
        raise ValueError("File type tidak didukung (hanya PDF, DOCX, TXT)")


# ==========================
# 🔹 EKSTRAK GAMBAR DARI FILE
# ==========================
def extract_images(file: UploadFile):
    filename = file.filename.lower()
    images = []

    # === PDF ===
    if filename.endswith(".pdf"):
        file.file.seek(0)
        with fitz.open(stream=file.file.read(), filetype="pdf") as pdf:
            for page in pdf:
                for img in page.get_images(full=True):
                    xref = img[0]
                    base_image = pdf.extract_image(xref)
                    image_bytes = base_image["image"]
                    image = Image.open(io.BytesIO(image_bytes)).convert("RGB")
                    images.append(image)

    # === DOCX ===
    elif filename.endswith(".docx"):
        file.file.seek(0)
        doc = Document(io.BytesIO(file.file.read()))
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                image_data = rel.target_part.blob
                image = Image.open(io.BytesIO(image_data)).convert("RGB")
                images.append(image)

    return images


# ==========================
# 🔹 EKSTRAK OCR DARI PDF
# ==========================
def extract_images_with_ocr(file: UploadFile):
    """Ekstrak semua gambar dari PDF dan konversi ke teks via OCR"""
    file.file.seek(0)
    text_from_images = []
    doc = fitz.open(stream=file.file.read(), filetype="pdf")

    for page_index in range(len(doc)):
        page = doc[page_index]
        images = page.get_images(full=True)
        for img_index, img in enumerate(images):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]

            # OCR ke teks
            image = Image.open(io.BytesIO(image_bytes))
            ocr_text = pytesseract.image_to_string(image, lang="eng+ind").strip()
            if ocr_text:
                text_from_images.append(ocr_text)

    return "\n".join(text_from_images)


def extract_text_from_pdf(file_path: str) -> str:
    """Ekstrak teks dari file PDF."""
    text = ""
    with fitz.open(file_path) as doc:
        for page in doc:
            text += page.get_text()
    return text